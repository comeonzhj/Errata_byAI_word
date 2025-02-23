import sys
import json
import os
import random
import string
from datetime import datetime
from pathlib import Path
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                           QHBoxLayout, QPushButton, QTextEdit, QFileDialog,
                           QProgressBar, QLabel, QDialog, QLineEdit, QMessageBox)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from docx import Document
from docx.shared import RGBColor
import requests

class ConfigDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle('模型配置')
        self.resize(500, 400)
        
        layout = QVBoxLayout()
        
        # API配置
        self.api_url = QLineEdit()
        self.api_url.setPlaceholderText('输入API URL,如https://api.deepseek.com/chat/completions')
        self.api_key = QLineEdit()
        self.api_key.setPlaceholderText('输入API Key，如 sk-...')
        self.model = QLineEdit()
        self.model.setPlaceholderText('输入模型名称,如 deepseek-chat')
        
        # System Prompt配置
        self.system_prompt = QTextEdit()
        self.system_prompt.setPlaceholderText('输入System Prompt')
        default_prompt = '''作为一个细致耐心的文字秘书，对下面的句子进行错别字检查，按如下结构以 JOSN 格式输出：
{
    "content_0":"原始句子",
    "wrong":true,//是否有需要被修正的错别字，布尔类型
    "annotation":"",//批注内容，string类型。如果wrong为true给出修正的解释；如果 wrong 字段为 false，则为空值
    "content_1":""//修改后的句子，string类型。如果wrong为false则留空
}'''
        self.system_prompt.setText(default_prompt)
        
        # 添加控件到布局
        layout.addWidget(QLabel('API URL:'))
        layout.addWidget(self.api_url)
        layout.addWidget(QLabel('API Key:'))
        layout.addWidget(self.api_key)
        layout.addWidget(QLabel('模型名称:'))
        layout.addWidget(self.model)
        layout.addWidget(QLabel('System Prompt:'))
        layout.addWidget(self.system_prompt)
        
        # 保存按钮
        save_btn = QPushButton('保存配置')
        save_btn.clicked.connect(self.accept)
        layout.addWidget(save_btn)
        
        self.setLayout(layout)

class ErrataCorrectionThread(QThread):
    progress_updated = pyqtSignal(int)
    log_updated = pyqtSignal(str)
    finished = pyqtSignal(str)
    error_occurred = pyqtSignal(str)

    def __init__(self, doc_path, api_config):
        super().__init__()
        self.doc_path = doc_path
        self.api_config = api_config
        self.is_running = True

    def run(self):
        try:
            doc = Document(self.doc_path)
            total_paragraphs = sum(1 for para in doc.paragraphs if para.text.strip())
            processed = 0

            # 创建输出文件名
            output_path = self._generate_output_filename(self.doc_path)
            
            for para in doc.paragraphs:
                if not self.is_running:
                    break
                    
                text = para.text.strip()
                if not text:
                    continue

                # 按句号分割文本
                sentences = text.split('。')
                
                # 清空段落内容，准备重新添加
                para.clear()
                
                for i, sentence in enumerate(sentences):
                    if not sentence.strip():
                        continue
                        
                    self.log_updated.emit(f'正在勘误：{sentence}')
                    
                    # 调用API进行勘误
                    result = self._check_sentence(sentence)
                    
                    # 添加原始句子
                    run = para.add_run(sentence)
                    
                    if result.get('wrong', False):
                        # 添加带高亮的批注
                        annotation_run = para.add_run(f' [修改建议：{result["annotation"]}]')
                        annotation_run.font.highlight_color = 7  # WD_COLOR_INDEX.YELLOW
                        
                        self.log_updated.emit(f'发现错误：{result["annotation"]}')
                        self.log_updated.emit('已在文档中添加修改建议')
                    else:
                        self.log_updated.emit('没有错误')
                    
                    # 如果不是最后一个句子，添加句号
                    if i < len(sentences) - 1 and sentence.strip():
                        para.add_run('。')

                processed += 1
                self.progress_updated.emit(int(processed * 100 / total_paragraphs))

            # 保存文档
            doc.save(output_path)
            self.finished.emit(output_path)
            
        except Exception as e:
            self.error_occurred.emit(str(e))

    def _check_sentence(self, sentence):
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {self.api_config["api_key"]}'
        }
        
        data = {
            'model': self.api_config['model'],
            'messages': [
                {'role': 'system', 'content': self.api_config['system_prompt']},
                {'role': 'user', 'content': sentence}
            ],
            'stream': False
        }
        
        response = requests.post(self.api_config['api_url'], headers=headers, json=data)
        response_data = response.json()
        
        try:
            result = json.loads(response_data['choices'][0]['message']['content'])
            return result
        except:
            return {'wrong': False, 'annotation': '', 'content_1': ''}

    def _generate_output_filename(self, original_path):
        dir_path = os.path.dirname(original_path)
        filename = os.path.basename(original_path)
        name, ext = os.path.splitext(filename)
        
        new_name = f'{name}_AI勘误{ext}'
        new_path = os.path.join(dir_path, new_name)
        
        # 如果文件已存在，添加随机后缀
        while os.path.exists(new_path):
            random_suffix = ''.join(random.choices(string.ascii_letters + string.digits, k=4))
            new_name = f'{name}_AI勘误_{random_suffix}{ext}'
            new_path = os.path.join(dir_path, new_name)
            
        return new_path

    def stop(self):
        self.is_running = False

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('AI勘误器_Word版')
        self.resize(800, 600)
        
        # 初始化UI
        self._init_ui()
        
        # 初始化变量
        self.current_file = None
        self.correction_thread = None
        self.api_config = {
            'api_url': '',
            'api_key': '',
            'model': '',
            'system_prompt': ''
        }
        
        # 加载配置
        self._load_config()

    def _init_ui(self):
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        
        # 顶部按钮区域
        button_layout = QHBoxLayout()
        self.open_btn = QPushButton('打开文档')
        self.config_btn = QPushButton('模型配置')
        self.start_btn = QPushButton('开始勘误')
        self.stop_btn = QPushButton('停止')
        self.export_btn = QPushButton('导出日志')
        
        self.open_btn.clicked.connect(self.open_file)
        self.config_btn.clicked.connect(self.show_config)
        self.start_btn.clicked.connect(self.start_correction)
        self.stop_btn.clicked.connect(self.stop_correction)
        self.export_btn.clicked.connect(self.export_log)
        
        button_layout.addWidget(self.open_btn)
        button_layout.addWidget(self.config_btn)
        button_layout.addWidget(self.start_btn)
        button_layout.addWidget(self.stop_btn)
        button_layout.addWidget(self.export_btn)
        
        # 进度条
        self.progress_bar = QProgressBar()
        
        # 日志区域
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        
        layout.addLayout(button_layout)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.log_text)
        
        # 初始化按钮状态
        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(False)
        self.export_btn.setEnabled(False)

    def open_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            '选择Word文档',
            '',
            'Word文档 (*.docx)'
        )
        
        if file_path:
            self.current_file = file_path
            self.start_btn.setEnabled(True)
            self.log_text.append(f'已选择文件：{file_path}')

    def _get_config_path(self):
        config_dir = Path.home() / '.errata_word'
        config_dir.mkdir(exist_ok=True)
        return config_dir / 'config.json'
    
    def _save_config(self):
        config_path = self._get_config_path()
        try:
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(self.api_config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            QMessageBox.warning(self, '警告', f'保存配置失败：{str(e)}')
    
    def _load_config(self):
        config_path = self._get_config_path()
        if config_path.exists():
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    self.api_config.update(json.load(f))
            except Exception as e:
                QMessageBox.warning(self, '警告', f'加载配置失败：{str(e)}')
    
    def show_config(self):
        dialog = ConfigDialog(self)
        
        # 填充现有配置
        dialog.api_url.setText(self.api_config['api_url'])
        dialog.api_key.setText(self.api_config['api_key'])
        dialog.model.setText(self.api_config['model'])
        if self.api_config['system_prompt']:
            dialog.system_prompt.setText(self.api_config['system_prompt'])
        
        if dialog.exec():
            self.api_config['api_url'] = dialog.api_url.text()
            self.api_config['api_key'] = dialog.api_key.text()
            self.api_config['model'] = dialog.model.text()
            self.api_config['system_prompt'] = dialog.system_prompt.toPlainText()
            self._save_config()
            self.log_text.append('配置已更新并保存')

    def start_correction(self):
        if not self.current_file:
            QMessageBox.warning(self, '警告', '请先选择Word文档')
            return
            
        if not all([self.api_config['api_url'], self.api_config['api_key'], 
                    self.api_config['model'], self.api_config['system_prompt']]):
            QMessageBox.warning(self, '警告', '请先完成模型配置')
            return
        
        self.correction_thread = ErrataCorrectionThread(self.current_file, self.api_config)
        self.correction_thread.progress_updated.connect(self.update_progress)
        self.correction_thread.log_updated.connect(self.update_log)
        self.correction_thread.finished.connect(self.correction_finished)
        self.correction_thread.error_occurred.connect(self.handle_error)
        
        self.correction_thread.start()
        
        # 更新按钮状态
        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.open_btn.setEnabled(False)
        self.config_btn.setEnabled(False)

    def stop_correction(self):
        if self.correction_thread and self.correction_thread.isRunning():
            self.correction_thread.stop()
            self.log_text.append('正在停止勘误...')

    def correction_finished(self, output_path):
        self.log_text.append(f'勘误完成！文件已保存至：{output_path}')
        self.reset_ui_state()
        self.export_btn.setEnabled(True)

    def handle_error(self, error_msg):
        QMessageBox.critical(self, '错误', f'发生错误：{error_msg}')
        self.reset_ui_state()

    def reset_ui_state(self):
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.open_btn.setEnabled(True)
        self.config_btn.setEnabled(True)
        self.progress_bar.setValue(0)

    def update_progress(self, value):
        self.progress_bar.setValue(value)

    def update_log(self, message):
        self.log_text.append(message)

    def export_log(self):
        options = ['导出全部', '仅导出错误']
        dialog = QDialog(self)
        dialog.setWindowTitle('导出日志')
        layout = QVBoxLayout(dialog)
        
        for option in options:
            btn = QPushButton(option)
            btn.clicked.connect(lambda checked, o=option: self._do_export_log(o, dialog))
            layout.addWidget(btn)
            
        dialog.exec()

    def _do_export_log(self, option, dialog):
        if not self.current_file:
            return
            
        log_content = self.log_text.toPlainText()
        if option == '仅导出错误':
            # 只保留包含"发现错误："的行
            lines = log_content.split('\n')
            log_content = '\n'.join(line for line in lines if '发现错误：' in line)
        
        # 生成日志文件名
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        log_filename = f'勘误日志_{timestamp}.txt'
        log_path = os.path.join(os.path.dirname(self.current_file), log_filename)
        
        try:
            with open(log_path, 'w', encoding='utf-8') as f:
                f.write(log_content)
            QMessageBox.information(self, '成功', f'日志已导出至：{log_path}')
        except Exception as e:
            QMessageBox.critical(self, '错误', f'导出日志失败：{str(e)}')
        
        dialog.close()

def main():
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())

if __name__ == '__main__':
    main()